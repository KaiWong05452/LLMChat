from docx import Document
import os
class QuestionExporter:
    def __init__(self, data, input_filename, output_path):
        self.data = data
        self.input_filename = input_filename
        self.output_path = output_path

        os.makedirs(self.output_path, exist_ok=True)

    def export_teacher_document(self):
        doc = Document()
        doc.add_heading(f"{self.input_filename} - Teacher Version", level=1)

        for i, question_data in enumerate(self.data, start=1):
            doc.add_heading(f"Question {i}:", level=2)
            doc.add_paragraph(question_data["Question"])

            doc.add_heading("Answer Choices:", level=3)
            for choice, answer in question_data["Answers"].items():
                doc.add_paragraph(f"{choice}. {answer}")

            doc.add_heading("Correct Answer:", level=3)
            doc.add_paragraph(f"{question_data['Correct_answer']}. {question_data['Answers'][question_data['Correct_answer']]}")

            doc.add_heading("Explanation:", level=3)
            doc.add_paragraph(question_data["Explanation"])

        doc.save(f"{self.output_path}/{self.input_filename}_teacherVersion.docx")

    def export_student_document(self):
        doc = Document()
        doc.add_heading(f"{self.input_filename} - Student Version", level=1)

        for i, question_data in enumerate(self.data, start=1):
            doc.add_heading(f"Question {i}:", level=2)
            doc.add_paragraph(question_data["Question"])

            doc.add_heading("Answer Choices:", level=3)
            for choice, answer in question_data["Answers"].items():
                doc.add_paragraph(f"{choice}. {answer}")

        doc.save(f"{self.output_path}/{self.input_filename}_studentVersion.docx")

